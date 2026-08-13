from docx import Document
from src.detectors.hybrid_detector import HybridDetector
from src.redactors.mapper import StatefulMapper

class DocxHandler:
    """
    Handles reading, redacting, and saving DOCX files safely without destroying
    document structure, tables, headers, or footers.
    """
    def __init__(self, detector: HybridDetector, mapper: StatefulMapper):
        self.detector = detector
        self.mapper = mapper

    def process_document(self, input_path: str, output_path: str):
        """
        Loads a document, redacts PII across all sections, and saves it.
        """
        doc = Document(input_path)

        # 1. Process Core Paragraphs
        self._process_paragraphs(doc.paragraphs)

        # 2. Process Tables
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    is_address_cell = False
                    if row_idx > 0:
                        try:
                            above_cell = table.cell(row_idx - 1, cell_idx)
                            if above_cell.text.strip().upper() in ["REGISTERED OFFICE", "CORPORATE OFFICE"]:
                                is_address_cell = True
                        except Exception:
                            pass
                            
                    self._process_paragraphs(cell.paragraphs, is_address_cell=is_address_cell)

        # 3. Process Headers and Footers
        for section in doc.sections:
            self._process_paragraphs(section.header.paragraphs)
            self._process_paragraphs(section.footer.paragraphs)

        doc.save(output_path)
        print(f"Redacted document successfully saved to {output_path}")

    def _process_paragraphs(self, paragraphs, is_address_cell=False):
        """
        Iterates over a list of paragraphs and redacts PII within them.
        """
        for paragraph in paragraphs:
            runs_text = "".join(r.text for r in paragraph.runs)
            if not runs_text.strip():
                continue

            if is_address_cell:
                entities = [{"entity_type": "ADDRESS", "start": 0, "end": len(runs_text), "text": runs_text}]
            else:
                # Detect entities in the run-reconstructed paragraph text
                entities = self.detector.analyze(runs_text)
            
            if not entities:
                continue

            # We process entities in reverse order (right to left) so that modifying
            # the text of a run does not invalidate the start/end character indices 
            # of entities that appear earlier in the paragraph.
            entities.sort(key=lambda x: x['start'], reverse=True)
            
            self._redact_runs(paragraph, entities)

    def _redact_runs(self, paragraph, entities):
        """
        Replaces detected PII across Word 'Runs' safely.
        A single paragraph is composed of multiple runs (sections of text with consistent formatting).
        An entity might span across multiple runs, so we must map character indices to runs.
        """
        for entity in entities:
            e_start = entity['start']
            e_end = entity['end']
            e_type = entity['entity_type']
            e_text = entity['text']

            fake_text = self.mapper.get_replacement(e_type, e_text)

            # Re-calculate run indices dynamically just to be safe, 
            # since we are modifying runs in-place.
            # Because we process right-to-left, the indices to the left remain intact.
            current_len = 0
            runs_info = []
            for idx, run in enumerate(paragraph.runs):
                runs_info.append({
                    "index": idx,
                    "run": run,
                    "start": current_len,
                    "end": current_len + len(run.text),
                    "text": run.text
                })
                current_len += len(run.text)

            # Find which runs the entity overlaps with
            overlapping_runs = []
            for r_info in runs_info:
                # Overlap condition:
                # Entity starts before run ends AND Entity ends after run starts
                if e_start < r_info['end'] and e_end > r_info['start']:
                    overlapping_runs.append(r_info)

            if not overlapping_runs:
                continue

            # If the entity is entirely contained in one run
            if len(overlapping_runs) == 1:
                r_info = overlapping_runs[0]
                run = r_info['run']
                local_start = e_start - r_info['start']
                local_end = e_end - r_info['start']
                
                # Replace the text inside the run
                run.text = run.text[:local_start] + fake_text + run.text[local_end:]
            
            # If the entity spans multiple runs
            else:
                first_r = overlapping_runs[0]
                last_r = overlapping_runs[-1]
                
                local_start = e_start - first_r['start']
                local_end = e_end - last_r['start']

                # Place the replacement text in the first run, keep the text before the entity
                first_r['run'].text = first_r['run'].text[:local_start] + fake_text
                
                # Clear all intermediate runs completely
                for r_info in overlapping_runs[1:-1]:
                    r_info['run'].text = ""
                
                # Keep the text after the entity in the last run
                last_r['run'].text = last_r['run'].text[local_end:]
